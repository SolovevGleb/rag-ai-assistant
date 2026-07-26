from docx import Document
from pathlib import Path


# где лежит исходный файл
input_file = Path("../raw/ElectricMotorController-RAG.docx")

# куда сохранить текст
output_file = Path("../processed/ElectricMotorController-RAG_text.txt")


def extract_text(docx_path):
    doc = Document(docx_path)

    result = []

    # достаём обычный текст
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text:
            result.append(text)

    # достаём таблицы
    for table in doc.tables:
        for row in table.rows:
            row_text = []

            for cell in row.cells:
                row_text.append(cell.text.strip())

            result.append(" | ".join(row_text))

    return "\n\n".join(result)


text = extract_text(input_file)


output_file.parent.mkdir(
    parents=True,
    exist_ok=True
)

with open(output_file, "w", encoding="utf-8") as f:
    f.write(text)


print("Готово")
print(f"Символов извлечено: {len(text)}")